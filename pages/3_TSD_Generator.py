import streamlit as st
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import tempfile
import datetime
import os
import subprocess  # Added for LibreOffice conversion
import pprint


# Function to convert DOCX to PDF using LibreOffice
def convert_docx_to_pdf_libreoffice(input_docx_path, output_dir):
    # Ensure the output directory exists
    os.makedirs(output_dir, exist_ok=True)

    # Construct the command to run LibreOffice in headless mode
    command = [
        "libreoffice",  # untuk versi Linux
        # "soffice",  # untuk versi Windows
        "--headless",
        "--convert-to", "pdf",
        input_docx_path,
        "--outdir", output_dir
    ]

    try:
        # Execute the command
        result = subprocess.run(command, check=True, capture_output=True, text=True)
        print(f"LibreOffice conversion stdout: {result.stdout}")
        if result.stderr:
            print(f"LibreOffice conversion stderr: {result.stderr}")

        # The output PDF will be in output_dir with the same base name as the DOCX
        pdf_filename = os.path.basename(input_docx_path).replace(".docx", ".pdf")
        output_pdf_path = os.path.join(output_dir, pdf_filename)

        if not os.path.exists(output_pdf_path):
            raise FileNotFoundError(f"LibreOffice did not produce the expected PDF at {output_pdf_path}")

        return output_pdf_path
    except subprocess.CalledProcessError as e:
        print(f"Error during LibreOffice conversion: {e.stderr}")
        raise RuntimeError(f"PDF conversion failed: {e.stderr}")
    except FileNotFoundError:
        raise RuntimeError("LibreOffice command not found. Make sure LibreOffice is installed and in your PATH.")


# App title and setup
st.set_page_config(page_title="TSD Generator", layout="centered")
st.title("📄 Technical Specification Document Generator")

# Initialize session state for API Specs and Dependencies
if "flow_diagram" not in st.session_state:
    st.session_state.flow_diagram = []

st.subheader("Informasi Umum")
group = st.text_input("Group")
project_name = st.text_input("Project Name")
date = st.date_input("Date", value=datetime.date.today())

# Initialize PIC dan service di session state
if "pic_list" not in st.session_state:
    st.session_state.pic_list = []
if "service_specifications" not in st.session_state:
    st.session_state.service_specifications = []
if "show_other_service" not in st.session_state:
    st.session_state.show_other_service = False
if "other_service" not in st.session_state:
    st.session_state.other_service = {"title": "", "subtitle": ""}

# Tombol tambah PIC
if st.button("➕ Tambah PIC"):
    st.session_state.pic_list.append({
        "pic_name": "",
        "services": []
    })

# Loop tiap PIC
for i, pic in enumerate(st.session_state.pic_list):
    with st.expander(f"PIC #{i+1}", expanded=True):
        pic_name = st.text_input(f"Nama PIC #{i+1}", value=pic["pic_name"], key=f"pic_name_{i}")
        st.session_state.pic_list[i]["pic_name"] = pic_name

        # Tombol tambah service description untuk PIC ini
        if st.button(f"➕ Tambah Service Description PIC #{i+1}", key=f"add_service_{i}"):
            st.session_state.pic_list[i]["services"].append("")

        # Input untuk tiap deskripsi service
        for j, service in enumerate(pic["services"]):
            service_desc = st.text_input(f"Service Description #{j+1} PIC #{i+1}", value=service, key=f"service_{i}_{j}")
            st.session_state.pic_list[i]["services"][j] = service_desc

        # Tombol hapus PIC
        if st.button(f"🗑️ Hapus PIC #{i+1}", key=f"del_pic_{i}"):
            del st.session_state.pic_list[i]
            st.rerun()


st.subheader("1. Introduction")
purpose = st.text_input("Purpose")
scope = st.text_input("Scope")
impact_analysis = st.text_input("Impact Analysis")

st.subheader("2. System Architecture")
overview = st.text_input("Overview")

# Tombol tambah flow diagram baru
if st.button("➕ Tambah Flow Diagram"):
    st.session_state.flow_diagram.append({
        "diagram_name": "",
        "pictures": []  # list untuk menyimpan beberapa gambar
    })

# Loop tiap flow diagram untuk input detailnya
for idx, flow in enumerate(st.session_state.flow_diagram):
    with st.expander(f"Flow Diagram #{idx + 1}", expanded=True):
        # Input nama diagram
        st.session_state.flow_diagram[idx]["diagram_name"] = st.text_input(
            f"Diagram Name #{idx + 1}",
            value=flow["diagram_name"],
            key=f"diagram_name_{idx}"
        )

        # Input gambar, bisa upload lebih dari 1 (gunakan file_uploader dengan accept_multiple_files=True)
        uploaded_files = st.file_uploader(
            f"Upload Images for Diagram #{idx + 1} (multiple allowed)",
            type=["png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key=f"upload_images_{idx}"
        )

        # Jika ada file baru yang diupload, tambahkan ke list pictures
        if uploaded_files:
            # Agar tidak hilang upload lama, gabungkan dengan gambar lama
            existing_pictures = st.session_state.flow_diagram[idx].get("pictures", [])
            # Simpan file dalam memori, misal dalam list berisi bytes dan nama file
            for file in uploaded_files:
                # Cek apakah file sudah ada berdasarkan nama, supaya gak duplikat
                if file.name not in [pic["name"] for pic in existing_pictures]:
                    existing_pictures.append({
                        "name": file.name,
                        "data": file.read()
                    })
            st.session_state.flow_diagram[idx]["pictures"] = existing_pictures

        # Tampilkan thumbnail gambar yang sudah diupload
        pictures = st.session_state.flow_diagram[idx].get("pictures", [])
        if pictures:
            st.write("Uploaded Images:")
            for pic in pictures:
                st.image(pic["data"], caption=pic["name"], width=150)

        # Tambahkan tombol untuk menghapus flow diagram jika perlu
        if st.button(f"🗑️ Hapus Flow Diagram #{idx + 1}", key=f"del_flow_{idx}"):
            del st.session_state.flow_diagram[idx]
            st.rerun()

st.subheader("3. Technical Stack")
software = st.text_input("Software")
database = st.text_input("Database")
caching = st.text_input("Caching")
logging = st.text_input("Logging")

st.subheader("4. Data Management")

DATA_TYPES = [
    "BFILE", "BINARY_DOUBLE", "BINARY_FLOAT", "BLOB", "CHAR", "CHARACTER", "CLOB", "DATE",
    "DEC", "DECIMAL", "DOUBLE PRECISION", "EXPRESSION", "FLOAT", "INT", "INTEGER", "INTERVAL",
    "MLSLABEL", "NATIONAL", "NCHAR", "NCLOB", "NUMBER", "NUMERIC", "NVARCHAR2", "RAW", "REAL",
    "REF", "ROWID", "SMALLINT", "SYS", "TIMESTAMP", "UROWID", "VARCHAR", "VARCHAR2", "LONG"
]

if st.button("➕ Tambah Service"):
    st.session_state.service_specifications.append({
        "name": "",
        "db_name": "",
        "table_name": "",
        "columns": []
    })

for i, service in enumerate(st.session_state.service_specifications):
    with st.expander(f"Service #{i+1}", expanded=True):
        name = st.text_input(f"Service Name #{i+1}", value=service["name"], key=f"name_{i}")
        db_name = st.text_input(f"Database Name #{i+1}", value=service["db_name"], key=f"db_name_{i}")
        table_name = st.text_input(f"Table Name #{i+1}", value=service["table_name"], key=f"table_name_{i}")

        st.session_state.service_specifications[i]["name"] = name
        st.session_state.service_specifications[i]["db_name"] = db_name
        st.session_state.service_specifications[i]["table_name"] = table_name

        if st.button(f"➕ Tambah Column Service #{i+1}", key=f"add_column_{i}"):
            st.session_state.service_specifications[i]["columns"].append({
                "row": "",
                "data_type": ""
            })

        for j, column in enumerate(service.get("columns", [])):
            row = st.text_input(f"Column No #{j+1} Service #{i+1}", value=column["row"], key=f"row_{i}_{j}")
            # data_type = st.text_input(f"Data Type #{j+1} Service #{i+1}", value=column["data_type"], key=f"data_type_{i}_{j}")

            data_type = st.selectbox(
                f"Data Type #{j+1} Service #{i+1}",
                DATA_TYPES,
                index=DATA_TYPES.index(column["data_type"].split('(')[0]) if column["data_type"] else 0,
                key=f"data_type_{i}_{j}"
            )

            digit = st.text_input(
                f"Digit (opsional) #{j+1} Service #{i+1}",
                value=column.get("digit", ""),
                key=f"digit_{i}_{j}"
            )

            final_type = f"{data_type}({digit})" if digit.strip() else data_type

            st.session_state.service_specifications[i]["columns"][j]["row"] = row
            st.session_state.service_specifications[i]["columns"][j]["data_type"] = final_type

        # Tombol hapus PIC
        if st.button(f"🗑️ Hapus Service #{i + 1}", key=f"del_flow_{i}"):
            del st.session_state.service_specifications[i]
            st.rerun()

# Tombol tambah Other Service
if not st.session_state.show_other_service:
    if st.button("➕ Tambah Other Service"):
        st.session_state.show_other_service = True
else:
    with st.expander(f"Other Services (Service #{len(st.session_state.service_specifications)+1})", expanded=True):
        st.session_state.other_title = "Other Services"
        other_title = st.text_input(
            "Judul Other Service",
            value=st.session_state.other_service.get("title", ""),
            key="other_title"
        )

        st.session_state.other_subtitle = "there is no database used for other services"
        other_subtitle = st.text_input(
            "Sub Judul Other Service",
            value=st.session_state.other_service.get("subtitle", ""),
            key="other_subtitle"
        )

        st.session_state.other_service["title"] =  f"4.{len(st.session_state.service_specifications)+1} {other_title}"
        st.session_state.other_service["subtitle"] = other_subtitle

        # Tombol hapus Other Service
        if st.button("🗑️ Hapus Other Service"):
            st.session_state.show_other_service = False
            st.session_state.other_service = {"title": "", "subtitle": ""}
            st.rerun()

st.subheader("5. Deployment Environtment")
environment = st.text_input("Environment")
configuration = st.text_input("Configuration")
deployment_proccess = st.text_input("Deployment Proccess")

st.subheader("6. Delivery Approval")
with st.expander("IT Developer", expanded=False):
    dev_name = st.text_input("Name")
    dev_nip = st.text_input("NIP")
    st.session_state.dev_status = "IT DEVELOPER"
    dev_status = st.text_input(
        "Developer Status",
        key="dev_status"
    )
with st.expander("MGR Developer", expanded=False):
    mgr_name = st.text_input("Name MGR")
    mgr_nip = st.text_input("NIP MGR")
    st.session_state.mgr_status = "IT DEVELOPER (MGR)"
    mgr_status = st.text_input(
        "MGR Status",
        key="mgr_status"
    )
with st.expander("Approval", expanded=False):
    approval_name = st.text_input("Name Approval")
    approval_nip = st.text_input("NIP Approval")
    st.session_state.approval_status = "Pgs. DEPT HEAD OF RE-TAIL CHANNEL & SER-VICES DELIVERY"
    approval_status = st.text_input(
        "Approval Status",
        key="approval_status"
    )

# Submit button moved to bottom
if st.button("🚀 Generate TSD"):
    doc = DocxTemplate("TSD-Automation.docx")

    flow_diagram_context = []
    for flow in st.session_state.flow_diagram:
        pictures = []
        for pic in flow["pictures"]:
            # Simpan gambar ke file temporer supaya bisa dibaca InlineImage
            temp_img_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            temp_img_file.write(pic["data"])
            temp_img_file.flush()
            temp_img_file.close()

            img = InlineImage(doc, temp_img_file.name, width=Mm(80))  # ukuran gambar 80 mm
            pictures.append(img)

        flow_diagram_context.append({
            "diagram_name": flow["diagram_name"],
            "pictures": pictures
        })

    pic_context = []

    for pic in st.session_state.pic_list:
        if not pic["pic_name"] or not pic["services"]:
            continue
        pic_context.append({
            "pic_name": pic["pic_name"],
            "services": pic["services"]
        })

    service_specifications_context = []

    for s in st.session_state.service_specifications:
        if not s.get("name"):
            continue

        columns = []
        for c in s.get("columns", []):
            if c.get("row") and c.get("data_type"):
                columns.append({
                    "row": c["row"],
                    "data_type": c["data_type"]
                })

        service_specifications_context.append({
            "name": s["name"],
            "db_name": s.get("db_name", ""),
            "table_name": s.get("table_name", ""),
            "columns": columns
        })

    context = {
        "group": group,
        "project_name": project_name,
        "date": str(date),
        "pic_list": pic_context,
        "purpose": purpose,
        "scope": scope,
        "impact_analysis": impact_analysis,
        "overview": overview,
        "software": software,
        "database": database,
        "caching": caching,
        "logging": logging,
        "environment": environment,
        "configuration": configuration,
        "deployment_proccess": deployment_proccess,
        "service_specifications": service_specifications_context,
        "other_service_title": st.session_state.other_service["title"],
        "other_service_subtitle": st.session_state.other_service["subtitle"],
        "dev_name": dev_name,
        "dev_nip": dev_nip,
        "dev_status": dev_status,
        "mgr_name": mgr_name,
        "mgr_nip": mgr_nip,
        "mgr_status": mgr_status,
        "approval_name": approval_name,
        "approval_nip": approval_nip,
        "approval_status": approval_status,
        "flow_diagram": flow_diagram_context,
    }

    # st.text("DEBUG: Context yang dikirim ke template:")
    # st.text(pprint.pformat(context))

    doc.render(context)

    # print(context["api_dependencies"])

    def remove_last_cell_except_first_last(doc):
        total_tables = len(doc.tables)

        for idx, table in enumerate(doc.tables):
            if idx == 0 or idx == total_tables - 1:
                continue

            last_row = table.rows[-1]
            # last_cell = last_row.cells[-1]
            # print(f"Menghapus isi cell terakhir: '{last_cell.text}'")
            tbl = table._tbl
            tbl.remove(last_row._tr)


    # def print_tables_last_row(doc):
    #     for i, table in enumerate(doc.tables):
    #         last_row = table.rows[-1]
    #         print(f"Table {i}, last row cells:")
    #         for j, cell in enumerate(last_row.cells):
    #             print(f" Cell {j}: '{cell.text}'")
    #
    # print_tables_last_row(doc)

    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    remove_last_cell_except_first_last(doc)
    doc.save(output_path)

    output_pdf = convert_docx_to_pdf_libreoffice(output_path, tempfile.gettempdir())

    with open(output_pdf, "rb") as f:
        st.success("✅ TSD berhasil dibuat!")
        st.download_button("⬇️ Download TSD", f, file_name="Generated_TSD.pdf")